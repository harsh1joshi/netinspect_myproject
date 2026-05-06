from flask import Flask, render_template, request, jsonify, redirect, session, Response, url_for
import sqlite3, json, time, hashlib, io, base64
from docx import Document
from docx.shared import Inches
from scanner import scan_target

app = Flask(__name__)
app.secret_key = "netinspect-secret"
DB_NAME = "netinspect.db"

# ---------------- DB ----------------
def get_db():
    conn = sqlite3.connect(DB_NAME)
    conn.row_factory = sqlite3.Row
    return conn

def hash_username(u):
    return hashlib.sha256(u.lower().encode()).hexdigest()

def get_scan(scan_id):
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT * FROM scans WHERE id=?", (scan_id,))
    row = c.fetchone()
    conn.close()
    return row

# ---------------- AUTH ----------------
@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form["username"]
        password = request.form["password"]

        if username == "harsh#2098" and password == "#@$&#q!":
            session["user_hash"] = hash_username(username)
            return redirect(url_for("index"))
        else:
            return render_template("login.html", error="Invalid username or password")

    return render_template("login.html")


@app.route("/logout")
def logout():
    session.clear()
    return redirect("/login")


@app.route("/")
def dashboard():
    if "user_hash" not in session:
        return redirect("/login")
    return render_template("index.html")


@app.route("/index")
def index():
    if "user_hash" not in session:
        return redirect("/login")
    return render_template("index.html")


# ---------------- SCAN ----------------
@app.route("/scan", methods=["POST"])
def scan():
    if "user_hash" not in session:
        return jsonify({"error": "Unauthorized"}), 401

    target = request.form["target"]
    mode = request.form["mode"]

    start = time.time()

    try:
        result = scan_target(target, mode)
    except Exception as e:
        return jsonify({"error": str(e)})

    elapsed = round(time.time() - start, 2)

    conn = get_db()
    c = conn.cursor()
    c.execute(
        "INSERT INTO scans(user_hash,target,mode,result,time_taken) VALUES(?,?,?,?,?)",
        (session["user_hash"], target, mode, json.dumps(result), elapsed)
    )
    scan_id = c.lastrowid
    conn.commit()
    conn.close()

    return jsonify({
        "status": "completed",
        "data": result,
        "time": elapsed,
        "scan_id": scan_id
    })


# ---------------- HISTORY ----------------
@app.route("/history")
def history():
    if "user_hash" not in session:
        return jsonify([])

    conn = get_db()
    c = conn.cursor()
    c.execute(
        "SELECT id,target,mode,time_taken,timestamp FROM scans ORDER BY timestamp DESC"
    )
    rows = c.fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])


@app.route("/scan/<int:scan_id>")
def open_old(scan_id):
    row = get_scan(scan_id)
    if not row:
        return ("Not found", 404)
    return jsonify(json.loads(row["result"]))


# ---------------- CSV EXPORT ----------------
@app.route("/export/csv/<int:scan_id>")
def export_csv(scan_id):
    row = get_scan(scan_id)
    if not row:
        return ("Not found", 404)

    data = json.loads(row["result"])
    if not data:
        return ("No data", 400)

    ports = data[0]["ports"]

    def generate():
        yield "Port,Service,Version,Finding,Severity\n"
        for p in ports:
            yield f'{p["port"]},{p["service"]},{p["version"]},"{p["vulnerability"]}",{p["severity"]}\n'

    return Response(
        generate(),
        mimetype="text/csv",
        headers={
            "Content-Disposition": f"attachment; filename=scan_{scan_id}.csv"
        }
    )


# ---------------- WORD EXPORT ----------------
@app.route("/export/word/<int:scan_id>", methods=["POST"])
def export_word(scan_id):
    row = get_scan(scan_id)
    if not row:
        return ("Not found", 404)

    chart = request.json["chart"]
    chart_bytes = chart.split(",", 1)[1]
    chart_bytes = io.BytesIO(base64.b64decode(chart_bytes))

    data = json.loads(row["result"])
    ports = data[0]["ports"] if data else []

    doc = Document()
    doc.add_heading(f"Security Scan Report – {row['target']}", 1)
    doc.add_paragraph(f"Mode: {row['mode']}")
    doc.add_paragraph(f"Time: {row['time_taken']}s")

    doc.add_picture(chart_bytes, width=Inches(4))

    table = doc.add_table(rows=1, cols=5)
    headers = ["Port","Service","Version","Finding","Severity"]

    for i,h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for p in ports:
        r = table.add_row().cells
        r[0].text = str(p["port"])
        r[1].text = p["service"]
        r[2].text = p["version"]
        r[3].text = p["vulnerability"]
        r[4].text = p["severity"]

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return Response(
        buffer,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename=scan_{scan_id}.docx"
        }
    )


if __name__ == "__main__":
    app.run(debug=True)